"""Rule-based ATS (Applicant Tracking System) compliance checks.

The rubric here is grounded in widely-published ATS guidance (Jobscan,
TopResume, and other resume-optimization sources): ATS parsers extract
plain text from a resume, so layout elements that don't survive that
extraction (tables, embedded images, header/footer content) silently
drop information, and keyword overlap with the job description is
consistently cited as the single biggest ranking factor.

Score weighting (100 pts total):
- Formatting & parseability   25 pts  (tables / images / header-footer)
- Sections & contact info     15 pts  (standard headings, email, phone)
- Content quality             20 pts  (quantified achievements, action verbs)
- Keyword match with the JD   40 pts  (computed by the caller from the
                                        Gemini-derived matched/unmatched
                                        skill lists)
"""

import io
import re

from docx import Document
from pypdf import PdfReader

from schemas import ATSCategory, ATSCheckItem, ATSReport

ACTION_VERBS = {
    "achieved", "administered", "analyzed", "architected", "automated",
    "built", "collaborated", "created", "delivered", "deployed",
    "designed", "developed", "directed", "drove", "engineered",
    "enhanced", "established", "executed", "generated", "implemented",
    "improved", "increased", "initiated", "launched", "led", "managed",
    "negotiated", "optimized", "orchestrated", "organized", "oversaw",
    "pioneered", "planned", "produced", "reduced", "refactored",
    "resolved", "scaled", "spearheaded", "streamlined", "strengthened",
    "supervised", "transformed",
}

SECTION_ALIASES = {
    "Experience": {
        "experience", "work experience", "professional experience",
        "employment history", "relevant experience",
    },
    "Education": {"education"},
    "Skills": {"skills", "technical skills", "core competencies"},
}

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[\s.-]?)?\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}")
BULLET_PREFIX_RE = re.compile(r"^\s*[-•*▪●◦‣∙·]\s+")
DIGIT_RE = re.compile(r"\d")


class RuleBasedFindings:
    def __init__(
        self,
        structure: ATSCategory,
        sections: ATSCategory,
        content: ATSCategory,
        missing_sections: list[str],
        quantified_ratio: float,
        action_verb_ratio: float,
        has_tables: bool,
        has_images: bool,
        has_header_footer_content: bool,
    ):
        self.structure = structure
        self.sections = sections
        self.content = content
        self.missing_sections = missing_sections
        self.quantified_ratio = quantified_ratio
        self.action_verb_ratio = action_verb_ratio
        self.has_tables = has_tables
        self.has_images = has_images
        self.has_header_footer_content = has_header_footer_content

    def summary_for_prompt(self) -> str:
        lines = [
            f"- Missing standard section headings: {', '.join(self.missing_sections) or 'none'}",
            f"- Bullets/lines with quantified results (numbers, %, $): {self.quantified_ratio:.0%}",
            f"- Bullets/lines starting with a strong action verb: {self.action_verb_ratio:.0%}",
            f"- Tables detected in document structure: {self.has_tables}",
            f"- Embedded images/graphics detected: {self.has_images}",
            f"- Contact info found only in header/footer (invisible to many ATS parsers): {self.has_header_footer_content}",
        ]
        return "\n".join(lines)


def _docx_structure_signals(content: bytes) -> tuple[bool, bool, bool]:
    document = Document(io.BytesIO(content))
    has_tables = len(document.tables) > 0
    has_images = len(document.inline_shapes) > 0
    header_footer_text = ""
    for section in document.sections:
        header_footer_text += " ".join(p.text for p in section.header.paragraphs)
        header_footer_text += " ".join(p.text for p in section.footer.paragraphs)
    has_header_footer_content = bool(header_footer_text.strip())
    return has_tables, has_images, has_header_footer_content


def _pdf_structure_signals(content: bytes) -> tuple[bool, bool, bool]:
    reader = PdfReader(io.BytesIO(content))
    has_images = any(len(page.images) > 0 for page in reader.pages)
    # Tables and header/footer regions aren't reliably separable from a raw
    # PDF text/image stream without layout analysis, so those checks are
    # skipped (treated as passing) for PDFs rather than guessed at.
    return False, has_images, False


def _structure_category(filename: str, content: bytes) -> tuple[ATSCategory, bool, bool, bool]:
    is_docx = filename.lower().endswith(".docx")
    if is_docx:
        has_tables, has_images, has_header_footer_content = _docx_structure_signals(content)
    else:
        has_tables, has_images, has_header_footer_content = _pdf_structure_signals(content)

    checks = [
        ATSCheckItem(
            label="No tables in layout",
            passed=not has_tables,
            detail=(
                "Tables found — ATS parsers often scramble or drop text inside table cells."
                if has_tables
                else "No tables detected."
                if is_docx
                else "Not reliably detectable for PDF; assumed pass."
            ),
        ),
        ATSCheckItem(
            label="No embedded images/graphics",
            passed=not has_images,
            detail=(
                "Images or graphics found — ATS systems can only read text, so this content is invisible to them."
                if has_images
                else "No embedded images detected."
            ),
        ),
        ATSCheckItem(
            label="Contact info not confined to header/footer",
            passed=not has_header_footer_content,
            detail=(
                "Text found in the document header/footer — many ATS parsers skip these regions entirely."
                if has_header_footer_content
                else "No header/footer content detected."
                if is_docx
                else "Not reliably detectable for PDF; assumed pass."
            ),
        ),
    ]
    points = [9, 8, 8]
    score = sum(p for item, p in zip(checks, points) if item.passed)
    category = ATSCategory(
        name="Formatting & Parseability", score=score, max_score=25, checks=checks
    )
    return category, has_tables, has_images, has_header_footer_content


def _sections_category(text: str) -> tuple[ATSCategory, list[str]]:
    # Some PDF layouts (columns/text boxes) collapse section headings into the
    # same line as surrounding content during text extraction, so headings
    # can't always be found by matching a line in isolation. Normalize
    # whitespace and search for the heading as a bounded phrase anywhere in
    # the text instead.
    normalized = re.sub(r"\s+", " ", text).strip().lower()

    email_found = bool(EMAIL_RE.search(text))
    phone_found = bool(PHONE_RE.search(text))

    section_present = {
        name: any(
            re.search(rf"\b{re.escape(alias)}\b", normalized) for alias in aliases
        )
        for name, aliases in SECTION_ALIASES.items()
    }
    missing_sections = [name for name, present in section_present.items() if not present]

    checks = [
        ATSCheckItem(
            label="Email address found",
            passed=email_found,
            detail="Found an email address." if email_found else "No email address detected in the resume text.",
        ),
        ATSCheckItem(
            label="Phone number found",
            passed=phone_found,
            detail="Found a phone number." if phone_found else "No phone number detected in the resume text.",
        ),
    ]
    for name, present in section_present.items():
        checks.append(
            ATSCheckItem(
                label=f'"{name}" section heading',
                passed=present,
                detail=(
                    f"Found a standard \"{name}\" heading."
                    if present
                    else f'No standard "{name}" heading found — ATS parsers rely on these to categorize content.'
                ),
            )
        )

    points = [3, 3, 3, 3, 3]
    score = sum(p for item, p in zip(checks, points) if item.passed)
    category = ATSCategory(name="Sections & Contact Info", score=score, max_score=15, checks=checks)
    return category, missing_sections


def _content_lines(text: str) -> list[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    content = []
    for line in lines:
        stripped = BULLET_PREFIX_RE.sub("", line).strip()
        if not stripped or len(stripped.split()) < 3:
            continue
        if EMAIL_RE.search(stripped) or PHONE_RE.search(stripped):
            continue
        normalized = stripped.rstrip(":").lower()
        if any(normalized in aliases for aliases in SECTION_ALIASES.values()):
            continue
        content.append(stripped)
    return content


def _content_category(text: str) -> tuple[ATSCategory, float, float]:
    lines = _content_lines(text)

    if lines:
        quantified_ratio = sum(1 for line in lines if DIGIT_RE.search(line)) / len(lines)
        action_verb_ratio = sum(
            1
            for line in lines
            if re.split(r"\W+", line)[0].lower() in ACTION_VERBS
        ) / len(lines)
    else:
        quantified_ratio = 0.0
        action_verb_ratio = 0.0

    quantified_points = round(min(1.0, quantified_ratio / 0.6) * 10)
    action_verb_points = round(min(1.0, action_verb_ratio / 0.5) * 10)

    checks = [
        ATSCheckItem(
            label="Quantified achievements",
            passed=quantified_ratio >= 0.6,
            detail=(
                f"{quantified_ratio:.0%} of content lines include a number, %, or $ "
                "(recruiter research targets 60%+)."
            ),
        ),
        ATSCheckItem(
            label="Action-verb-led bullets",
            passed=action_verb_ratio >= 0.5,
            detail=(
                f"{action_verb_ratio:.0%} of content lines start with a strong action verb "
                "(e.g. Built, Led, Optimized)."
            ),
        ),
    ]
    category = ATSCategory(
        name="Content Quality",
        score=quantified_points + action_verb_points,
        max_score=20,
        checks=checks,
    )
    return category, quantified_ratio, action_verb_ratio


def gather_findings(filename: str, content: bytes, resume_text: str) -> RuleBasedFindings:
    structure, has_tables, has_images, has_header_footer_content = _structure_category(
        filename, content
    )
    sections, missing_sections = _sections_category(resume_text)
    content_category, quantified_ratio, action_verb_ratio = _content_category(resume_text)

    return RuleBasedFindings(
        structure=structure,
        sections=sections,
        content=content_category,
        missing_sections=missing_sections,
        quantified_ratio=quantified_ratio,
        action_verb_ratio=action_verb_ratio,
        has_tables=has_tables,
        has_images=has_images,
        has_header_footer_content=has_header_footer_content,
    )


def build_report(
    findings: RuleBasedFindings, matched_skills: list[str], unmatched_skills: list[str]
) -> ATSReport:
    total_skills = len(matched_skills) + len(unmatched_skills)
    match_rate = len(matched_skills) / total_skills if total_skills else 0.0
    keyword_score = round(match_rate * 40)

    keyword_category = ATSCategory(
        name="Keyword Match with Job Description",
        score=keyword_score,
        max_score=40,
        checks=[
            ATSCheckItem(
                label="Job description keyword coverage",
                passed=match_rate >= 0.75,
                detail=(
                    f"Resume evidences {len(matched_skills)} of {total_skills} "
                    f"job-description keywords/skills ({match_rate:.0%} match rate; "
                    "ATS-optimization tools generally target 75%+)."
                ),
            )
        ],
    )

    categories = [findings.structure, findings.sections, findings.content, keyword_category]
    overall_score = sum(category.score for category in categories)

    return ATSReport(overall_score=overall_score, categories=categories)
